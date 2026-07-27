"""Build publication tables directly from the audited analysis outputs."""

from __future__ import annotations

import csv
import math
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "08_Analysis_Outputs" / "tables"
MAIN_TABLES = ROOT / "02_Main_Tables"
SUPPLEMENT_TABLES = ROOT / "04_Supplementary_Materials"


def rows(filename: str) -> list[dict[str, str]]:
    with (OUTPUT / filename).open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def number(value: str | float, digits: int = 2) -> str:
    return f"{float(value):.{digits}f}"


def integer(value: str | float) -> str:
    return f"{int(float(value)):,}"


def p_value(value: str | float | None) -> str:
    if value in (None, "", "NA", "NaN"):
        return "-"
    value = float(value)
    if math.isnan(value):
        return "-"
    if value < 0.001:
        return "<0.001"
    return f"{value:.3f}"


def is_missing(value: str | float | None) -> bool:
    return value is None or str(value).strip() in ("", "NA", "NaN", "nan")


def estimate(row: dict[str, str], prefix: str, digits: int = 2) -> str:
    lower_key = prefix + "_lower"
    upper_key = prefix + "_upper"
    if lower_key not in row and prefix.endswith("_HR"):
        lower_key = prefix[:-3] + "_lower"
        upper_key = prefix[:-3] + "_upper"
    if lower_key not in row:
        lower_key, upper_key = "lower", "upper"
    return (
        f"{float(row[prefix]):.{digits}f} "
        f"({float(row[lower_key]):.{digits}f}-"
        f"{float(row[upper_key]):.{digits}f})"
    )


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 90,
                     bottom: int = 80, end: int = 90) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (
        ("top", top), ("start", start), ("bottom", bottom), ("end", end)
    ):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def new_document(landscape: bool = False) -> Document:
    document = Document()
    section = document.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    return document


def title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(8)


def subtitle(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(75, 75, 75)


def make_table(document: Document, headers: list[str], data: list[list[str]],
               font_size: float = 8, first_col_left: bool = True):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    header = table.rows[0]
    set_repeat_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, "D9E2F3")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Arial"
        run.font.size = Pt(font_size)
    for row_values in data:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            set_cell_margins(cells[index])
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if first_col_left and index == 0
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
    return table


def save(document: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def add_table1(document: Document) -> None:
    title(document, "Table 1. Baseline characteristics of the study population, by urinary albumin-to-creatinine ratio (UACR) category")
    headers = [
        "Characteristic", "Overall", "Normal UACR\n(<10 mg/g)",
        "Low-grade albuminuria\n(10-29 mg/g)",
        "Albuminuria\n(>=30 mg/g)", "P value", "Max SMD"
    ]
    table = make_table(document, headers, [], font_size=7.2)
    current_section = None
    for row in rows("Table1_baseline.csv"):
        if row["section"] != current_section:
            current_section = row["section"]
            cells = table.add_row().cells
            merged = cells[0]
            for cell in cells[1:]:
                merged = merged.merge(cell)
            set_cell_shading(merged, "EFEFEF")
            run = merged.paragraphs[0].add_run(current_section)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(7.5)
        cells = table.add_row().cells
        values = [
            row["characteristic"], row["overall"].strip(), row["normal"].strip(),
            row["subclinical"].strip(), row["albuminuria"].strip(),
            p_value(row["p"]),
            "-" if not row["max_pairwise_smd"] else number(row["max_pairwise_smd"], 3),
        ]
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(7.2)
            if row["row_type"] == "category" and index == 0:
                run.bold = True
    note(
        document,
        "Values are survey-weighted mean (SD), median [IQR], or weighted percentage; N is "
        "unweighted. The weighted population estimate represents the average annual US "
        "non-institutionalized adult population over the pooled 1999-2018 survey period, not a "
        "count of distinct persons. P values account for the NHANES design. Max SMD is the largest pairwise "
        "standardized mean difference across UACR groups. Abbreviations: BP, blood pressure; "
        "Prescription-recorded aspirin or antiplatelet use is derived from the NHANES prescription-medication record and does not capture unrecorded over-the-counter aspirin use. "
        "eGFR, estimated glomerular filtration rate; HbA1c, glycated hemoglobin; HDL, "
        "high-density lipoprotein; IQR, interquartile range; SMD, standardized mean difference; "
        "UACR, urinary albumin-to-creatinine ratio."
    )


def add_table2(document: Document) -> None:
    title(document, "Table 2. Adjusted hazard ratios for all-cause mortality by UACR category, across sequential adjustment models (Models 1-5)")
    data = []
    for row in rows("model_hierarchy_complete_case.csv"):
        data.append([
            row["model"], integer(row["N"]), integer(row["events"]),
            "1.00 (reference)", row["subclinical"], row["subclinical_p"],
            row["albuminuria"], row["albuminuria_p"],
        ])
    make_table(
        document,
        ["Model", "N", "Deaths", "Normal UACR\n(<10 mg/g)",
         "Low-grade HR (95% CI)", "P",
         "Albuminuria HR (95% CI)", "P"],
        data, font_size=7.8,
    )
    note(
        document,
        "Survey-weighted complete-case Cox models; normal UACR (<10 mg/g) is the reference. "
        "M1 unadjusted; M2 age, sex, race/ethnicity, cycle; M3 adds BMI, smoking, diabetes, "
        "BP-treatment category, eGFR; M4 adds PIR and education; M5 adds dyslipidemia. "
        "Abbreviations: BMI, body-mass index; BP, blood pressure; CI, confidence interval; "
        "eGFR, estimated glomerular filtration rate; HR, hazard ratio; PIR, poverty-income "
        "ratio; UACR, urinary albumin-to-creatinine ratio."
    )


MAIN_BUILDERS = [add_table1, add_table2]
def _bold(document, text):
    p = document.add_paragraph(text)
    p.runs[0].bold = True


def _incidence_body(document):
    _bold(document, "Panel A. All-cause mortality incidence")
    incidence = []
    for row in rows("S2_incidence.csv"):
        incidence.append([
            row["uacr_category"], integer(row["N"]), integer(row["deaths"]),
            integer(round(float(row["person_years_unweighted"]))),
            f"{float(row['incidence_rate_per_1000_py']):.1f} "
            f"({float(row['rate_lower']):.1f}-{float(row['rate_upper']):.1f})",
        ])
    make_table(document, ["UACR category", "N", "Deaths", "Person-years",
                          "Rate per 1,000 person-years (95% CI)"], incidence)
    _bold(document, "Panel B. Marginally standardized risk")
    risk = []
    for row in rows("S2_absolute_risk.csv"):
        ard = "Reference" if is_missing(row["ARD_lower"]) else (
            f"{float(row['ARD_percentage_points']):.1f} "
            f"({float(row['ARD_lower']):.1f}-{float(row['ARD_upper']):.1f})")
        risk.append([integer(row["horizon_years"]), row["uacr_category"],
                     f"{float(row['adjusted_risk_percent']):.1f} "
                     f"({float(row['risk_lower']):.1f}-{float(row['risk_upper']):.1f})", ard])
    make_table(document, ["Horizon, y", "UACR category", "Adjusted risk, % (95% CI)",
                          "Risk difference, percentage points (95% CI)"], risk)


def sup_s1(document: Document) -> None:
    title(document, "Supplementary Table S1. Mortality incidence rates and marginally standardized 10-year absolute mortality risk, by UACR category")
    _incidence_body(document)
    note(document, "Panel A rates are design-based survey ratios. Panel B uses complete-case Model 3, a tie-corrected weighted Breslow baseline hazard, and marginal standardization; CIs are from a stratified delete-1 PSU jackknife (JKn), in which every replicate refits the model and re-estimates the baseline hazard.")


def sup_s2(document: Document) -> None:
    title(document, "Supplementary Table S3. Primary complete-case and multiple-imputation hazard ratios for all-cause mortality (Models 1-5)")
    data = []
    for filename in ("model_hierarchy_complete_case.csv", "model_hierarchy_mi.csv"):
        for row in rows(filename):
            data.append([row["analysis"], row["model"], integer(row["N"]), integer(row["events"]),
                         row["subclinical"], row["subclinical_p"], row["albuminuria"], row["albuminuria_p"]])
    make_table(document, ["Analysis", "Model", "N", "Deaths", "Low-grade HR (95% CI)", "P",
                          "Albuminuria HR (95% CI)", "P"], data, font_size=7.5)
    note(document, "Complete-case analysis is primary. Multiple imputation by chained equations (m=20; 10 iterations) was a sensitivity analysis. Body-mass index, smoking, family income-to-poverty ratio, and education were imputed using predictive mean matching. The imputation model included all Model 5 covariates, the event indicator, follow-up time, the Nelson-Aalen cumulative hazard, sampling weight, stratum, and PSU. Exposure, outcome, and survey-design variables were retained as observed; estimates were combined using Rubin rules.")


def sup_s3(document: Document) -> None:
    title(document, "Supplementary Table S4. Continuous-UACR analyses: spline functional form and measurement-error (disattenuation) sensitivity")
    _bold(document, "Functional form")
    row = rows("S4_functional_form.csv")[0]
    data = [["HR per doubling of UACR",
             f"{float(row['HR_per_doubling']):.2f} ({float(row['lower']):.2f}-{float(row['upper']):.2f})"],
            ["P, overall spline association", p_value(row["p_overall_spline"])],
            ["P, nonlinearity", p_value(row["p_nonlinearity"])],
            ["Knot locations, mg/g", ", ".join(number(row[f"knot_{i}_mg_g"], 1) for i in range(1, 5))]]
    make_table(document, ["Quantity", "Value"], data)
    _bold(document, "Knot sensitivity")
    knot_data = [[r["specification"], r["knot_locations_mg_g"], p_value(r["overall_p"]), p_value(r["nonlinearity_p"])]
                 for r in rows("S4_knot_sensitivity.csv")]
    make_table(document, ["Specification", "Knot locations, mg/g", "Overall P", "Nonlinearity P"], knot_data, font_size=7.8)
    _bold(document, "Measurement-error sensitivity")
    me = []
    for r in rows("S5_measurement_error.csv"):
        label = "Uncorrected" if float(r["reliability"]) == 1 else number(r["reliability"], 2)
        me.append([label, f"{float(r['corrected_HR']):.2f} ({float(r['lower']):.2f}-{float(r['upper']):.2f})"])
    make_table(document, ["Assumed reliability", "HR per doubling (95% CI)"], me)
    note(document, "Survey-weighted complete-case Model 3; restricted cubic spline on log2(UACR) with knots at the 5th, 35th, 65th, and 95th percentiles. The measurement-error analysis divides the log-HR and its standard error by an assumed reliability (hypothetical, not estimated from repeat samples); this is a deterministic disattenuation, not regression calibration.")


def sup_s4(document: Document) -> None:
    title(document, "Supplementary Table S5. Sensitivity analyses of the UACR-all-cause-mortality association (definitional, subgroup-stratum, inflammatory-marker, covariate-omission, and E-value analyses)")
    data = []
    for row in rows("S8_sensitivity_complete_case.csv"):
        data.append([row["analysis"], integer(row["N"]), integer(row["events"]),
                     estimate(row, "subclinical_HR"), estimate(row, "albuminuria_HR")])
    mi = next(r for r in rows("model_hierarchy_mi.csv") if r["model"] == "M3")
    data.append(["Multiple imputation, Model 3", integer(mi["N"]), integer(mi["events"]),
                 mi["subclinical"], mi["albuminuria"]])
    fasting_labels = {
        "Fasting-subsample Model 3: primary definitions": "Fasting subsample, Model 3, primary diabetes definition",
        "Fasting-subsample Model 3: fasting-expanded definitions": "Fasting subsample, Model 3, fasting-expanded diabetes definition",
        "Fasting-subsample Model 5: primary definitions": "Fasting subsample, Model 5, primary diabetes and dyslipidemia definitions",
        "Fasting-subsample Model 5: fasting-expanded definitions": "Fasting subsample, Model 5, fasting-expanded diabetes and dyslipidemia definitions",
    }
    for row in rows("S20_fasting_subsample_sensitivity.csv"):
        data.append([fasting_labels[row["analysis"]], integer(row["N"]), integer(row["events"]),
                     estimate(row, "subclinical_HR"), estimate(row, "albuminuria_HR")])
    egfr = rows("S9_egfr_interaction.csv")
    for row in egfr:
        data.append([row["analysis"].replace("eGFR ", "eGFR stratum ") + " mL/min/1.73 m2",
                     integer(row["N"]), integer(row["events"]),
                     estimate(row, "subclinical_HR"), estimate(row, "albuminuria_HR")])
    for row in rows("S15_mortality_crp.csv"):
        data.append([f"{row['assay']}, {row['model']}", integer(row["N"]), integer(row["deaths"]),
                     estimate(row, "subclinical_HR"), estimate(row, "albuminuria_HR")])
    make_table(document, ["Analysis", "N", "Deaths", "Low-grade HR (95% CI)", "Albuminuria HR (95% CI)"],
               data, font_size=7.6)
    _bold(document, "Covariate-omission robustness")
    cov = [[r["model"], integer(r["N"]), integer(r["deaths"]),
            estimate(r, "subclinical_HR"), estimate(r, "albuminuria_HR")]
           for r in rows("S16_covariate_omission.csv")
           if r["block"] == "Covariate-omission robustness"]
    make_table(document, ["Model / omitted covariate", "N", "Deaths",
                          "Low-grade HR (95% CI)", "Albuminuria HR (95% CI)"], cov, font_size=7.5)
    _bold(document, "E-values for the primary Model 3 estimates")
    evalues = [[r["exposure"], number(r["E_value_point"], 2), number(r["E_value_confidence_limit"], 2)]
               for r in rows("S8_evalues.csv")]
    make_table(document, ["Exposure", "E-value, point estimate", "E-value, confidence limit"], evalues)
    note(document, f"UACR indicates urinary albumin-to-creatinine ratio; HR, hazard ratio; CI, confidence interval; eGFR, estimated glomerular filtration rate; CRP, C-reactive protein; LDL, low-density lipoprotein. Normal UACR (<10 mg/g) is the reference; all models are survey-weighted, complete-case Model 3 unless otherwise stated. eGFR-stratum models omit continuous eGFR (design-based eGFR-by-UACR interaction P={float(egfr[0]['interaction_p']):.3f}). CRP-adjustment rows are restricted to each CRP assay block and are not pooled. Covariate-omission rows sequentially remove blocks of Model 3 covariates to assess dependence on any single adjustment. Fasting-subsample models use fasting weights (WTSAF4YR/5 for 1999-2002; WTSAF2YR/10 for 2003-2018); the fasting-expanded diabetes definition adds fasting glucose >=126 mg/dL, and the fasting-expanded dyslipidemia definition adds LDL cholesterol >=160 mg/dL or fasting triglycerides >=200 mg/dL. E-values quantify the minimum strength of association, on the risk-ratio scale, that an unmeasured confounder would need with both exposure and mortality to fully explain the observed association.")


def sup_s5(document: Document) -> None:
    title(document, "Supplementary Table S6. Expanded low-range UACR categories and all-cause mortality")
    source = rows("S18_low_range_uacr.csv")
    data = []
    for row in source:
        hr = ("1.00 (reference)" if is_missing(row["lower"]) else
              f"{float(row['HR']):.2f} ({float(row['lower']):.2f}-{float(row['upper']):.2f})")
        data.append([row["uacr_category"], integer(row["N"]), integer(row["deaths"]), hr, p_value(row["p"])])
    make_table(document, ["UACR category", "N", "Deaths", "HR (95% CI)", "P"], data)
    note(document, f"Exploratory survey-weighted complete-case Model 3; <5 mg/g is reference. P for linear trend {p_value(source[0]['p_trend'])}. Categories are not proposed as diagnostic or treatment thresholds.")


def sup_s6(document: Document) -> None:
    title(document, "Supplementary Table S7. Proportional-hazards diagnostics (per covariate) and baseline-stratified sensitivity analysis")
    labels = {"uacr_cat": "UACR category", "age": "Age", "sex_female": "Sex",
              "race_eth": "Race or ethnicity", "cycle": "NHANES cycle", "bmi": "Body-mass index",
              "smoking": "Smoking status", "dm_main": "Diabetes", "bp_stage": "BP-treatment group",
              "egfr": "eGFR", "GLOBAL": "Global test"}
    _bold(document, "Scaled Schoenfeld-residual tests, per covariate")
    data = [[labels.get(r["term"], r["term"]), number(r["chisq"], 2), integer(r["df"]), p_value(r["p"])]
            for r in rows("S6_ph_diagnostics.csv")]
    make_table(document, ["Term", "Chi-square", "df", "P"], data)
    _bold(document, "Sensitivity: UACR hazard ratios with the baseline stratified by nonproportional covariates")
    sens = [[r["model"], r["subclinical"], r["albuminuria"]] for r in rows("S6_ph_sensitivity.csv")]
    make_table(document, ["Model", "Low-grade HR (95% CI)", "Albuminuria HR (95% CI)"], sens, font_size=7.8)
    note(document, "UACR indicates urinary albumin-to-creatinine ratio; HR, hazard ratio; CI, confidence interval; BP, blood pressure; eGFR, estimated glomerular filtration rate. Scaled Schoenfeld-residual tests are from the unweighted complete-case Model 3. UACR category itself satisfies the proportional-hazards assumption (P=.18), whereas age, sex, body-mass index, smoking, BP-treatment group, and eGFR show evidence of nonproportionality (all P<.05; age, BP-treatment group, and eGFR at P<.001). Because the baseline hazard can be stratified only by categorical variables, the sensitivity models stratify by the nonproportional categorical covariates (sex, smoking, BP-treatment group) and, in the final model, additionally by categorized age and eGFR while retaining their continuous forms for adjustment; body-mass index (the weakest violation) is retained as a covariate. The UACR hazard ratios are essentially unchanged across all specifications. Survey-weighted Cox estimation is used for the effect estimates reported elsewhere.")


def sup_s7(document: Document) -> None:
    title(document, "Supplementary Table S8. Incremental predictive value of UACR for 10-year mortality: discrimination (Harrell C) and categorical reclassification (NRI)")
    _bold(document, "Discrimination (Harrell C-statistic)")
    data = []
    for row in rows("S7_discrimination.csv"):
        cstat = f"{float(row['C_statistic']):.4f} ({float(row['lower']):.4f}-{float(row['upper']):.4f})"
        change = "-" if is_missing(row["change_in_C"]) else (
            f"{float(row['change_in_C']):+.4f} ({float(row['change_lower']):+.4f} to {float(row['change_upper']):+.4f})")
        data.append([row["model"], cstat, change])
    make_table(document, ["Model", "Harrell C-statistic (95% CI)", "Change in C (95% CI)"], data)
    _bold(document, "Categorical net reclassification improvement")
    nri = [[r["component"], f"{float(r['NRI']):+.3f} ({float(r['lower']):+.3f} to {float(r['upper']):+.3f})"]
           for r in rows("S19_ipcw_nri.csv")]
    make_table(document, ["Component", "NRI (95% CI)"], nri)
    note(document, "Exploratory 10-year incremental value of adding UACR to the clinical model; NRI risk bands are <5%, 5%-<10%, 10%-<20%, and >=20%. Point estimates use pooled MEC sampling weights (NRI also uses inverse-probability-of-censoring weights). CIs use a stratified delete-1 PSU jackknife (JKn) with 153 design degrees of freedom; every replicate refits the model(s).")


def sup_s8(document: Document) -> None:
    title(document, "Supplementary Table S9. Heart-disease mortality: cause-specific and competing-risk hazard ratios and endpoint-definition sensitivity")
    _bold(document, "Model hierarchy (cause-specific Cox)")
    data = []
    for row in rows("S10_heart_disease_hierarchy.csv"):
        data.append([row["model"], integer(row["N"]), integer(row["events"]),
                     estimate(row, "subclinical_HR"), p_value(row["subclinical_p"]),
                     estimate(row, "albuminuria_HR"), p_value(row["albuminuria_p"])])
    make_table(document, ["Model", "N", "Heart-disease deaths", "Low-grade HR (95% CI)", "P",
                          "Albuminuria HR (95% CI)", "P"], data, font_size=7.8)
    _bold(document, "Competing-risk analysis (Fine-Gray)")
    cr = [[r["analysis"], estimate(r, "subclinical_HR"), estimate(r, "albuminuria_HR")]
          for r in rows("S11_competing_risk.csv")]
    make_table(document, ["Analysis", "Low-grade HR (95% CI)", "Albuminuria HR (95% CI)"], cr)
    _bold(document, "Endpoint-definition sensitivity (1999-2014)")
    ep = [[r["endpoint"], r["cycles"], integer(r["events"]),
           estimate(r, "subclinical_HR"), estimate(r, "albuminuria_HR")]
          for r in rows("S17_endpoint_sensitivity.csv")]
    make_table(document, ["Endpoint definition", "Cycles", "Events", "Low-grade HR (95% CI)", "Albuminuria HR (95% CI)"], ep)
    note(document, "Cause-specific survey-weighted Cox models censor non-heart-disease deaths; Fine-Gray models treat non-heart-disease death as the competing event. Heart-disease death uses NCHS underlying leading-cause code 001. The endpoint-definition sensitivity is restricted to 1999-2014 because cerebrovascular code 005 is unavailable in the 2015-2018 public-use files.")


def sup_s9(document: Document) -> None:
    title(document, "Supplementary Table S10. Adjusted association of UACR with C-reactive protein (exploratory)")
    data = []
    for row in rows("S14_log_crp_models.csv"):
        data.append([row["assay"], row["contrast"], estimate(row, "geometric_mean_ratio"),
                     p_value(row["p"]), p_value(row["uacr_by_bp_interaction_p"])])
    make_table(document, ["Assay", "Contrast", "Geometric-mean ratio (95% CI)", "P", "UACR x BP interaction P"], data)
    avail = {r["assay"]: integer(r["N"]) for r in rows("S12_crp_availability.csv")}
    note(document, f"UACR indicates urinary albumin-to-creatinine ratio; CRP, C-reactive protein; BP, blood pressure; CI, confidence interval. Exploratory survey-weighted linear models of log CRP, adjusted as in Model 3; ratios compare each UACR category with normal UACR (<10 mg/g). Standard and high-sensitivity assays are analyzed separately and are not pooled: standard CRP was measured in 1999-2010 (unweighted N={avail.get('standard CRP','-')}) and high-sensitivity CRP in 2015-2018 (N={avail.get('hs-CRP','-')}); CRP was not measured in 2011-2014. The geometric-mean CRP by UACR within each BP-treatment category is shown in Supplementary Figure S9.")


def sup_joint(document: Document) -> None:
    title(document, "Supplementary Table S2. Standardized 10-year mortality risk and hazard ratios across the 18 joint blood-pressure-treatment and UACR groups")
    short = {"Normal UACR (<10 mg/g)": "Normal (<10)",
             "Low-grade albuminuria (10-29 mg/g)": "Low-grade (10-29)",
             "Albuminuria (>=30 mg/g)": "Albuminuria (>=30)"}
    _bold(document, "Panel A. Standardized 10-year risk and hazard ratio versus the global reference (non-elevated BP with normal UACR)")
    data = []
    for r in rows("joint_18_cells.csv"):
        risk = f"{float(r['adjusted_10y_risk_percent']):.1f} ({float(r['risk_lower']):.1f}-{float(r['risk_upper']):.1f})"
        if r["bp_category"].startswith("Non-elevated") and r["uacr_category"].startswith("Normal"):
            hr = "1.00 (reference)"
        else:
            hr = f"{float(r['HR_vs_global_reference']):.2f} ({float(r['HR_lower']):.2f}-{float(r['HR_upper']):.2f})"
        data.append([r["bp_category"], short.get(r["uacr_category"], r["uacr_category"]),
                     integer(r["N"]), integer(r["deaths"]), risk, hr])
    make_table(document, ["BP-treatment group", "UACR", "N", "Deaths",
                          "Standardized 10-y risk, % (95% CI)", "HR vs global reference (95% CI)"],
               data, font_size=7.5)
    _bold(document, "Panel B. Within-BP-group contrasts (each elevated UACR category vs normal UACR in the same BP-treatment group)")
    document.paragraphs[-1].paragraph_format.page_break_before = True
    wdata = []
    for r in rows("within_bp_contrasts.csv"):
        hr = f"{float(r['HR']):.2f} ({float(r['HR_lower']):.2f}-{float(r['HR_upper']):.2f})"
        ard = f"{float(r['ARD_percentage_points']):.1f} ({float(r['ARD_lower']):.1f} to {float(r['ARD_upper']):.1f})"
        contrast = r["comparison"].replace("Low-grade albuminuria (10-29 mg/g)", "Low-grade").replace(
            "Albuminuria (>=30 mg/g)", "Albuminuria").replace("Normal UACR (<10 mg/g)", "normal")
        wdata.append([r["bp_category"], contrast, hr, ard])
    make_table(document, ["BP-treatment group", "Contrast", "Within-group HR (95% CI)",
                          "Within-group ARD, percentage points (95% CI)"], wdata, font_size=7.5)
    note(document, "UACR indicates urinary albumin-to-creatinine ratio; BP, blood pressure; HR, hazard ratio; CI, confidence interval; ARD, absolute risk difference. Estimates use the complete-case survey-weighted 18-group Model 3. Risks are marginally standardized; global HRs use non-elevated BP with normal UACR as the reference, and Panel B compares UACR categories within BP groups. JKn CIs refit the model and baseline hazard; N and deaths are unweighted. Interaction P values were .781 for the omnibus three-level test (10 df), .715 for low-grade versus normal (5 df), and .809 for albuminuria versus normal (5 df).")


SUPPLEMENT_BUILDERS = [sup_s1, sup_joint, sup_s2, sup_s3, sup_s4,
                       sup_s5, sup_s6, sup_s7, sup_s8, sup_s9]


def build_individual() -> None:
    main_names = ["Table_1_Baseline.docx", "Table_2_AllCause_Model_Hierarchy.docx"]
    for builder, filename in zip(MAIN_BUILDERS, main_names):
        document = new_document(landscape=True)
        builder(document)
        save(document, MAIN_TABLES / filename)
    supplement_names = [
        "Supplementary_Table_S1_Incidence_AbsoluteRisk.docx",
        "Supplementary_Table_S2_Joint_18Group_Risk.docx",
        "Supplementary_Table_S3_CompleteCase_vs_MI.docx",
        "Supplementary_Table_S4_ContinuousUACR_Analyses.docx",
        "Supplementary_Table_S5_Sensitivity_Analyses.docx",
        "Supplementary_Table_S6_ExpandedLowRange_UACR.docx",
        "Supplementary_Table_S7_ProportionalHazards.docx",
        "Supplementary_Table_S8_IncrementalValue.docx",
        "Supplementary_Table_S9_HeartDisease_Analyses.docx",
        "Supplementary_Table_S10_CRP_Association.docx",
    ]
    for index, (builder, filename) in enumerate(
        zip(SUPPLEMENT_BUILDERS, supplement_names), start=1
    ):
        document = new_document(landscape=index in (2, 3, 5, 9))
        builder(document)
        save(document, SUPPLEMENT_TABLES / filename)


def build_merged() -> None:
    main = new_document(landscape=True)
    for index, builder in enumerate(MAIN_BUILDERS):
        if index:
            main.add_page_break()
        builder(main)
    save(main, MAIN_TABLES / "_MERGED_Main_Tables.docx")

    supplement = new_document(landscape=True)
    title(supplement, "Supplementary Tables")
    subtitle(
        supplement,
        "NHANES 1999-2018; complete-case primary analysis; normal UACR (<10 mg/g) "
        "is the reference unless stated otherwise."
    )
    for index, builder in enumerate(SUPPLEMENT_BUILDERS):
        if index:
            supplement.add_page_break()
        builder(supplement)
    save(supplement, SUPPLEMENT_TABLES / "_MERGED_Supplementary_Tables.docx")


if __name__ == "__main__":
    build_individual()
    build_merged()
    print("Main and supplementary tables completed.")
