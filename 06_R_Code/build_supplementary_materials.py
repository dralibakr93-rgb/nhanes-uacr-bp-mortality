"""Assemble ONE Supplementary Materials document: a clickable contents page with
page numbers, then all supplementary tables (S1-S10) and all supplementary
figures (S1-S10). Contents entries hyperlink to the item and show its page
(a PAGEREF field Word fills in on open). Reuses build_docx_tables' builders."""

from pathlib import Path
import sys
import build_docx_tables as B
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
FIGDIR = ROOT / "08_Analysis_Outputs" / "figures"
HIGHLIGHT_REVISIONS = "--highlight-revisions" in sys.argv
OUT_NAME = ("Supplementary_Materials_REVISED_HIGHLIGHTED.docx"
            if HIGHLIGHT_REVISIONS else "Supplementary_Materials_COMBINED.docx")
OUT = ROOT / "04_Supplementary_Materials" / OUT_NAME

MANUSCRIPT_TITLE = ("Low-Grade Albuminuria and Mortality Across Blood Pressure Levels and "
                    "Treatment Status in US Adults With Preserved eGFR and Without "
                    "Self-Reported Major Cardiovascular Disease")

TABLE_TITLES = {
 1: "Mortality incidence rates and marginally standardized 10-year absolute mortality risk, by UACR category",
 2: "Standardized 10-year mortality risk and hazard ratios across the 18 joint blood-pressure-treatment and UACR groups",
 3: "Primary complete-case and multiple-imputation hazard ratios for all-cause mortality (Models 1-5)",
 4: "Continuous-UACR analyses: spline functional form and measurement-error (disattenuation) sensitivity",
 5: "Sensitivity analyses of the UACR-all-cause-mortality association",
 6: "Expanded low-range UACR categories and all-cause mortality",
 7: "Proportional-hazards diagnostics (per covariate) and baseline-stratified sensitivity analysis",
 8: "Incremental predictive value of UACR for 10-year mortality: discrimination and reclassification",
 9: "Heart-disease mortality: cause-specific, competing-risk, and endpoint-definition analyses",
 10: "Adjusted association of UACR with C-reactive protein (exploratory)",
}
# (filename, contents-page title, body caption). KM figures (S4/S5/S8) carry
# their title on the image, so their body caption describes the figure without
# repeating that title (avoids the title appearing twice).
FIGS = {
 1: ("Supplementary_Figure_S1_Participant_Attrition.png",
     "Participant-level covariate missingness",
     "Participant-level covariate missingness. Number of participants missing each covariate in the calibrated analytic cohort (N=29,951); fasting-subsample and cycle-limited biomarkers (fasting lipids and glucose, C-reactive protein, complete blood count) account for most missingness."),
 2: ("Supplementary_Figure_S2_Missingness_byCycle.png",
     "Per-variable and per-cycle missingness",
     "Per-variable and per-cycle missingness. Percentage of participants missing each variable within each 2-year NHANES cycle, sorted by overall missingness."),
 3: ("Supplementary_Figure_S3_UACR_Prevalence.png",
     "Survey-weighted UACR-category prevalence by blood-pressure-treatment category",
     "Survey-weighted UACR-category prevalence within each blood-pressure-treatment category."),
4: ("Supplementary_Figure_S4_Landmark_KM_AllCause.png",
     "Two-year landmark all-cause mortality by UACR category",
     "Cumulative all-cause mortality by baseline UACR category. The first two years are shown as pre-landmark follow-up (light gray shading). At the dashed two-year landmark, follow-up is restarted at zero among participants alive and uncensored at two years. Curves are sampling-weighted, unadjusted Kaplan-Meier estimates, shaded bands show pointwise 95% confidence intervals, and the number-at-risk table reports unweighted participants at 0, 2, 4, 8, 12, and 16 years. Annotated hazard ratios are from the adjusted survey-weighted Cox landmark model among two-year survivors."),
 5: ("Supplementary_Figure_S5_Landmark_BPxUACR_AllCause.png",
     "Two-year landmark all-cause mortality by UACR within blood-pressure-treatment category",
     "Cumulative all-cause mortality by baseline UACR category within each of the six blood-pressure-treatment categories (panels A-F). In each panel, the first two years are shown as pre-landmark follow-up (light gray shading). At the dashed two-year landmark, follow-up is restarted at zero among participants alive and uncensored at two years. Curves are sampling weighted and unadjusted, shaded bands show pointwise 95% confidence intervals, and number-at-risk tables report unweighted participants at 0, 2, 4, 8, 12, and 16 years. Per-panel hazard ratios are from adjusted survey-weighted Cox landmark models among two-year survivors."),
 6: ("Supplementary_Figure_S6_Forest_Albuminuria_AllCause.png",
     "Subgroup forest plot of albuminuria and all-cause mortality",
     "Subgroup forest plot of the adjusted albuminuria (>=30 mg/g) versus normal-UACR hazard ratio for all-cause mortality."),
 7: ("Supplementary_Figure_S7_RCS_Continuous_HeartDisease.png",
     "Continuous UACR and heart-disease mortality (restricted cubic spline)",
     "Continuous UACR and heart-disease mortality modeled with a survey-weighted restricted cubic spline (reference, 10 mg/g)."),
 8: ("Supplementary_Figure_S8_Landmark_KM_HeartDisease.png",
     "Two-year landmark heart-disease mortality by UACR category",
     "Cause-specific heart-disease mortality by baseline UACR category (secondary/exploratory). The first two years are shown as pre-landmark follow-up (light gray shading). At the dashed two-year landmark, follow-up is restarted at zero among participants alive and uncensored at two years. Curves are sampling-weighted cause-specific Kaplan-Meier failure estimates with pointwise 95% confidence intervals; non-heart-disease deaths are censored, so the curves may exceed the competing-risk cumulative incidence. The number-at-risk table reports unweighted participants at 0, 2, 4, 8, 12, and 16 years. Annotated hazard ratios are from the adjusted survey-weighted Cox landmark model among two-year survivors."),
 9: ("Supplementary_Figure_S9_CRP_by_UACR.png",
     "Adjusted geometric-mean C-reactive protein by UACR (exploratory)",
     "Exploratory adjusted geometric-mean C-reactive protein across UACR categories, shown within each blood-pressure-treatment group (small multiples; standard-CRP block, 1999-2010). Points are geometric means with 95% confidence intervals."),
 10: ("Supplementary_Figure_S10_RiskProfile_BPxUACR.png",
     "Standardized 10-year mortality risk profiles by UACR across blood-pressure-treatment categories",
     "Standardized 10-year all-cause mortality risk profiles by UACR across the six blood-pressure-treatment categories."),
}

doc = B.new_document(landscape=False)
_id = [10]

# Landscape cover/contents page; tables and figures remain portrait.
cover_section = doc.sections[0]
cover_section.orientation = WD_ORIENT.LANDSCAPE
cover_section.page_width = Inches(11)
cover_section.page_height = Inches(8.5)
cover_section.top_margin = Inches(0.45)
cover_section.bottom_margin = Inches(0.45)
cover_section.left_margin = Inches(0.60)
cover_section.right_margin = Inches(0.60)

def bookmark_para(paragraph, name):
    _id[0] += 1
    s = OxmlElement('w:bookmarkStart'); s.set(qn('w:id'), str(_id[0])); s.set(qn('w:name'), name)
    e = OxmlElement('w:bookmarkEnd'); e.set(qn('w:id'), str(_id[0]))
    paragraph._p.insert(0, s); paragraph._p.append(e)

def _sized_rPr(*extra, bold=False, color=None, underline=False, highlight=False, sz='21'):
    rPr = OxmlElement('w:rPr')
    if bold: rPr.append(OxmlElement('w:b'))
    if color:
        c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
    if underline:
        u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    if highlight:
        h = OxmlElement('w:highlight'); h.set(qn('w:val'), 'yellow'); rPr.append(h)
    s = OxmlElement('w:sz'); s.set(qn('w:val'), sz); rPr.append(s)
    return rPr

def _field(parent, instr):
    """Append a Word field (begin/instr/sep/placeholder/end) to parent element."""
    for kind, txt in (('begin', None), ('instr', instr), ('sep', None), ('text', '1'), ('end', None)):
        r = OxmlElement('w:r'); r.append(_sized_rPr())
        if kind == 'instr':
            it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = txt; r.append(it)
        elif kind == 'text':
            t = OxmlElement('w:t'); t.text = txt; r.append(t)
        else:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), kind if kind != 'sep' else 'separate'); r.append(fc)
        parent.append(r)

def toc_entry(container, prefix, title, anchor, highlight=False):
    p = container.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs'); tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right'); tab.set(qn('w:leader'), 'dot'); tab.set(qn('w:pos'), '6200')
    tabs.append(tab); pPr.append(tabs)
    hl = OxmlElement('w:hyperlink'); hl.set(qn('w:anchor'), anchor)
    def run(text, bold):
        r = OxmlElement('w:r'); r.append(_sized_rPr(
            bold=bold, color='1F4E79', underline=True,
            highlight=highlight, sz='18'))
        t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; r.append(t)
        return r
    hl.append(run(prefix, True)); hl.append(run(title, False))
    p._p.append(hl)
    rt = OxmlElement('w:r'); rt.append(OxmlElement('w:tab')); p._p.append(rt)
    _field(p._p, f' PAGEREF {anchor} \\h ')
    return p

def h(text, size, center=False):
    p = doc.add_paragraph()
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
    return p

# ---- landscape contents page (tables and figures in balanced columns) ----
h("Supplementary Materials", 19, center=True)
_sub = doc.add_paragraph(); _sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
_r = _sub.add_run(MANUSCRIPT_TITLE); _r.italic = True; _r.font.size = Pt(10)
_r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
_sub.paragraph_format.space_after = Pt(10)

def section_header(container, text):
    if hasattr(container, 'paragraphs') and len(container.paragraphs) == 1 and not container.paragraphs[0].text:
        p = container.paragraphs[0]
    else:
        p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs'); tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right'); tab.set(qn('w:pos'), '6200'); tabs.append(tab); pPr.append(tabs)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    r2 = p.add_run("\tPage"); r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    # bottom border under the header for a clean divider
    pbdr = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6'); bot.set(qn('w:space'), '2'); bot.set(qn('w:color'), 'BBBBBB')
    pbdr.append(bot); pPr.append(pbdr)

def no_table_borders(table):
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:val'), 'nil')
        borders.append(node)
    tblPr.append(borders)

def _shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def _minimal_table_borders(table):
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for edge, val, size, color in (
        ('top', 'single', '8', '7F7F7F'),
        ('left', 'nil', '0', 'FFFFFF'),
        ('bottom', 'single', '8', '7F7F7F'),
        ('right', 'nil', '0', 'FFFFFF'),
        ('insideH', 'nil', '0', 'FFFFFF'),
        ('insideV', 'nil', '0', 'FFFFFF'),
    ):
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:val'), val)
        node.set(qn('w:sz'), size)
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), color)
        borders.append(node)
    tblPr.append(borders)

def style_supplementary_tables(document):
    """Gray headers and minimal rules, matching the earlier supplement style."""
    for table in document.tables:
        if (len(table.rows) == 1 and len(table.columns) == 2 and
                'Supplementary Table S1.' in table.cell(0, 0).text):
            no_table_borders(table)
            continue
        _minimal_table_borders(table)
        if not table.rows:
            continue
        for cell in table.rows[0].cells:
            _shade_cell(cell, 'BFBFBF')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)

        headers = [cell.text.strip() for cell in table.rows[0].cells]
        is_joint_panel_a = (
            len(headers) >= 6 and
            headers[0] == 'BP-treatment group' and
            headers[1] == 'UACR' and
            'Standardized 10-y risk' in headers[4]
        )
        if is_joint_panel_a and len(table.rows) >= 19:
            # Reproduce the prior three gray UACR blocks. Show the category label
            # once at the beginning of each six-row BP block.
            for start in (1, 7, 13):
                for row_index in range(start, start + 6):
                    _shade_cell(table.rows[row_index].cells[1], 'D9D9D9')
                    if row_index > start:
                        table.rows[row_index].cells[1].text = ''

toc_grid = doc.add_table(rows=1, cols=2)
toc_grid.autofit = False
no_table_borders(toc_grid)
left, right = toc_grid.rows[0].cells
left.width = Inches(4.75)
right.width = Inches(4.75)

section_header(left, "Supplementary Tables")
toc_entry(left, "Supplementary Methods. ", "Laboratory assay harmonization",
          "supp_methods", highlight=HIGHLIGHT_REVISIONS)
for n, t in TABLE_TITLES.items():
    toc_entry(left, f"Supplementary Table S{n}. ", t, f"tbl_S{n}")

section_header(right, "Supplementary Figures")
for n, (_, toc_title, _) in FIGS.items():
    toc_entry(right, f"Supplementary Figure S{n}. ", toc_title, f"fig_S{n}")

content_section = doc.add_section(WD_SECTION.NEW_PAGE)
content_section.orientation = WD_ORIENT.PORTRAIT
content_section.page_width = Inches(8.5)
content_section.page_height = Inches(11)

# ---- supplementary methods ----
methods_heading = doc.add_paragraph()
methods_heading.paragraph_format.space_after = Pt(6)
heading_run = methods_heading.add_run("Supplementary Methods. Laboratory Assay Harmonization")
heading_run.bold = True
heading_run.font.size = Pt(12)
if HIGHLIGHT_REVISIONS:
    heading_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
bookmark_para(methods_heading, "supp_methods")

methods_paragraph = doc.add_paragraph()
methods_paragraph.paragraph_format.space_after = Pt(8)
methods_text = (
    "Because serum and urinary creatinine assay methods changed across NHANES cycles, "
    "measurements were harmonized before calculating estimated glomerular filtration rate "
    "and urinary albumin-to-creatinine ratio. Serum creatinine from 1999-2000 was standardized "
    "as 0.147 + 1.013 × measured serum creatinine (mg/dL) before application of the 2021 "
    "CKD-EPI equation [38]. For urinary creatinine X (mg/dL) measured by the Jaffé method in "
    "1999-2006, conversion to the later enzymatic scale was applied as "
    "[1.02 × √X - 0.36]² when X <75, [1.05 × √X - 0.74]² when 75 ≤ X <250, and "
    "[1.01 × √X - 0.10]² when X ≥250 [39]."
)
methods_run = methods_paragraph.add_run(methods_text)
methods_run.font.size = Pt(10)
if HIGHLIGHT_REVISIONS:
    methods_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

# ---- tables (bookmark each title paragraph) ----
for n, builder in enumerate(B.SUPPLEMENT_BUILDERS, 1):
    builder(doc)
    for para in reversed(doc.paragraphs):
        if para.text.strip().startswith(f"Supplementary Table S{n}."):
            para.paragraph_format.page_break_before = True
            bookmark_para(para, f"tbl_S{n}"); break

style_supplementary_tables(doc)

# ---- figures (bookmark each caption) ----
for n, (fname, _toc_title, legend) in FIGS.items():
    cap = doc.add_paragraph()
    cap.paragraph_format.page_break_before = True
    cap.add_run(f"Supplementary Figure S{n}. ").bold = True
    cap.add_run(legend)
    bookmark_para(cap, f"fig_S{n}")
    path = FIGDIR / fname
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[missing figure file: {fname}]")

# ---- page-number footer on every section ("Page X") ----
for section in doc.sections:
    section.footer.is_linked_to_previous = False
    fp = section.footer.paragraphs[0]
    fp.text = ''
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead = fp.add_run("Page "); lead.font.size = Pt(9); lead.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    _field(fp._p, ' PAGE ')

# make Word refresh the PAGE / PAGEREF fields on open
uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true')
doc.settings.element.insert(0, uf)

doc.save(str(OUT))
print("Saved:", OUT.name, "| tables:", len(B.SUPPLEMENT_BUILDERS), "figures:", len(FIGS))
